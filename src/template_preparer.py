"""
Module for preparing and populating templates with Unit of Competency data
"""
import os
import shutil
import tempfile

from docx import Document
from docxtpl import DocxTemplate

from .utils import FileNames, Paths, StopExecution, unit_path_from_code


class TemplatePreparer:
    """Class for preparing and populating templates with UoC data"""
    
    def __init__(self, unit_code: str, uoc_data: dict, additional_details: dict = None):
        """
        Initialize template preparer
        
        Args:
            unit_code: The unit code
            uoc_data: Unit of Competency data
            additional_details: Additional template details (version, course info etc)
        """
        self.unit_code = unit_code
        self.unit_path = unit_path_from_code(unit_code)
        self.uoc_data = uoc_data
        self.template_context = {**uoc_data}
        
        if additional_details:
            self.template_context.update(additional_details)
    
    def _setup_unit_directory(self):
        """Create unit directory if it doesn't exist"""
        os.makedirs(self.unit_path, exist_ok=True)
    
    def _prepare_assessment_mapping(self):
        """Prepare the Assessment Mapping document"""
        jinja_template_path = os.path.join(Paths.Jinja_Templates, FileNames.Assessment_Mapping)
        output_path = os.path.join(self.unit_path, FileNames.Assessment_Mapping)
        
        # Create temp file for intermediate processing
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            # Copy template to temp file
            shutil.copy2(jinja_template_path, temp_file.name)
            
            # Render template with context
            template = DocxTemplate(temp_file.name)
            template.render(self.template_context)
            template.save(temp_file.name)
            
            # Further processing with python-docx
            doc = Document(temp_file.name)
            
            # Process Elements & Performance Criteria table
            table = doc.tables[1]
            self._clear_table_rows(table)
            row = 2
            for element in self.uoc_data["elements"]:
                start_row = row
                for pc in element["performance_criteria"]:
                    new_row = table.add_row()
                    new_row.cells[1].text = f"{pc['index']} {pc['description']}"
                
                table.rows[start_row].cells[0].text = f"{element['index']}. {element['title']}"
                if start_row != row:
                    self._merge_cells(table, start_row, row, 0)
                row += len(element["performance_criteria"])
            
            # Process Foundation Skills table
            table = doc.tables[2]
            self._clear_table_rows(table)
            row = 2
            for skill in self.uoc_data["foundational_skills"]:
                start_row = row
                for desc in skill["descriptions"]:
                    new_row = table.add_row()
                    new_row.cells[1].text = desc
                
                table.rows[start_row].cells[0].text = skill["skill"]
                if start_row != row:
                    self._merge_cells(table, start_row, row, 0)
                row += len(skill["descriptions"])
            
            # Process Evidence tables
            self._populate_evidence_table(doc.tables[4], self.uoc_data["performance_evidence"])
            self._populate_evidence_table(doc.tables[5], self.uoc_data["knowledge_evidence"])
            self._populate_evidence_table(doc.tables[6], self.uoc_data["assessment_conditions"])
            
            # Remove Range of Conditions table
            doc.tables[3]._tbl.getparent().remove(doc.tables[3]._tbl)
            
            # Save final document
            doc.save(output_path)
            
            # Cleanup temp file
            try:
                os.remove(temp_file.name)
            except Exception:
                pass  # Ignore cleanup errors
    
    @staticmethod
    def _clear_table_rows(table):
        """Remove all rows from table except headers"""
        for row in range(len(table.rows) - 1, 1, -1):
            table._tbl.remove(table.rows[row]._tr)
    
    @staticmethod
    def _merge_cells(table, start_row, end_row, col):
        """Merge cells in a column between start and end rows"""
        start_cell = table.rows[start_row].cells[col]
        end_cell = table.rows[end_row].cells[col]
        start_cell.merge(end_cell)
    
    @staticmethod
    def _populate_evidence_table(table, items):
        """Populate a simple evidence table with items"""
        TemplatePreparer._clear_table_rows(table)
        for item in items:
            new_row = table.add_row()
            new_row.cells[0].text = item
    
    def prepare_all_templates(self):
        """Prepare all templates for the unit"""
        self._setup_unit_directory()
        self._prepare_assessment_mapping()
        # Add preparation of other templates here as needed
